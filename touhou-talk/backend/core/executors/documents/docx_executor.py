# core/input_pipeline/executors/documents/docx_executor.py
"""
docx_executor.py
=================================================
Word (.docx) ファイル用 Executor（テキスト抽出版）

【設計思想】
- docx は「論理構造をある程度保持した文書」
- ここでは「段落」「見出し」「表」を最低限テキスト化する
- 見た目・装飾・画像・コメント・脚注は扱わない

【前提ライブラリ】
- python-docx
  pip install python-docx

【やらないこと】
- スタイル完全再現
- 表構造の厳密解析
- 画像抽出
- OCR
- マクロ処理

【やること】
- 段落テキスト抽出
- 見出しレベルのラベル化
- 表のセル内容をテキストに展開
- 安全なサイズ制限
"""

from __future__ import annotations

from typing import Any, Dict, Optional, List

from .base import DocumentExecutorBase

try:
    from docx import Document
except Exception:
    Document = None


class DOCXExecutor(DocumentExecutorBase):
    """
    DOCX 用 Executor
    """

    supported_extensions = (".docx",)
    filetype_label = "docx"

    # ===== 安全制御 =====
    MAX_BYTES: int = 10 * 1024 * 1024   # 10MB
    MAX_CHARS: int = 120_000            # 最大文字数
    MAX_PARAGRAPHS: int = 3_000         # 最大段落数
    MAX_TABLE_ROWS: int = 1_000         # 表の最大行数（全表合算）

    def _extract_text(
        self,
        *,
        filename: str,
        raw_bytes: bytes,
        encoding: Optional[str],
    ) -> tuple[str, Dict[str, Any]]:
        """
        DOCX からテキストを抽出する。

        戦略：
        1) ライブラリ存在確認
        2) サイズ制限（bytes）
        3) Document 初期化
        4) 段落抽出（見出し含む）
        5) 表抽出
        6) 正規化
        7) 長さ制限
        """

        if Document is None:
            raise RuntimeError(
                "python-docx is not installed. "
                "Install with: pip install python-docx"
            )

        if raw_bytes is None:
            raise ValueError("raw_bytes is None")

        original_size = len(raw_bytes)

        # 1) サイズ制限（bytes）
        truncated_bytes = raw_bytes
        truncated_by_bytes = False
        if original_size > self.MAX_BYTES:
            truncated_bytes = raw_bytes[: self.MAX_BYTES]
            truncated_by_bytes = True

        # 2) Document 初期化
        try:
            from io import BytesIO
            doc = Document(BytesIO(truncated_bytes))
        except Exception as e:
            raise ValueError(f"docx_open_failed: {e}")

        lines: List[str] = []
        paragraph_count = 0
        table_row_count = 0

        # 3) 段落抽出
        for p in doc.paragraphs:
            if paragraph_count >= self.MAX_PARAGRAPHS:
                break

            text = (p.text or "").strip()
            if not text:
                continue

            style_name = ""
            try:
                if p.style and p.style.name:
                    style_name = p.style.name
            except Exception:
                style_name = ""

            # 見出し判定（厳密でなくてよい）
            if style_name.lower().startswith("heading"):
                lines.append(f"[{style_name}] {text}")
            else:
                lines.append(text)

            paragraph_count += 1

        # 4) 表抽出（後段にまとめて出す）
        for table in doc.tables:
            for row in table.rows:
                if table_row_count >= self.MAX_TABLE_ROWS:
                    break

                cell_texts = []
                for cell in row.cells:
                    cell_text = (cell.text or "").strip()
                    if cell_text:
                        cell_texts.append(cell_text)

                if cell_texts:
                    lines.append("[TABLE] " + " | ".join(cell_texts))

                table_row_count += 1

            if table_row_count >= self.MAX_TABLE_ROWS:
                break

        # 5) 正規化
        combined = "\n".join(lines)
        normalized = (
            combined
            .replace("\r\n", "\n")
            .replace("\r", "\n")
        )

        # 6) 文字数制限
        truncated_by_chars = False
        if len(normalized) > self.MAX_CHARS:
            normalized = normalized[: self.MAX_CHARS]
            truncated_by_chars = True

        meta: Dict[str, Any] = {
            "original_bytes": original_size,
            "truncated_by_bytes": truncated_by_bytes,
            "truncated_by_chars": truncated_by_chars,
            "char_count": len(normalized),
            "paragraphs_read": paragraph_count,
            "max_paragraphs": self.MAX_PARAGRAPHS,
            "table_rows_read": table_row_count,
            "max_table_rows": self.MAX_TABLE_ROWS,
            "docx_text_extracted": True,
            "ocr_used": False,
        }

        return normalized, meta